#!/usr/bin/env python
# coding: utf-8

# In[6]:


from threading import Thread
from tkinter.font import Font
from tkinter.ttk import *
from tkinter.scrolledtext import ScrolledText
from tkinter.filedialog import askopenfile
from tkinter import *
from tkinter import messagebox
from functools import partial
from tkinter.messagebox import showinfo
from PIL import ImageTk
from docx import Document
from shutil import copy
from string import digits
import time
import datetime
import pandas
#import openpyxl as xl
import nltk
import re
import colorama
import numpy as np
import docx
from docx import Document
from nltk.tokenize import word_tokenize
from nltk.tokenize.treebank import TreebankWordDetokenizer
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from nltk.corpus import wordnet


lemmatizer = WordNetLemmatizer()
nltk.download('omw-1.4')
nltk.download('wordnet')
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('averaged_perceptron_tagger')


level_from_style_name = {f'Heading {i}': i for i in range(10)}
level_from_style_name2 = {f'Heading {i}': i for i in range(10)}
adjective_label, synonym_label, synonym_view, adjective_view, NoneType = None, None, None, None, None
vertical, horizontal, ambiguous_label = None, None, None
or_label, username_label, username_entry, password_label, password_entry, login_button, upload_menu = None, None, None, None, None, None, None
srs_document_button, uploaded_srs_label, uploaded_user_label, upload_srs_document_label, user_document_button = None, None, None, None, None
upload_user_document_label, upload_file_button, open_button = None, None, None
password_entry_2, password_label_2, sign_up_button, review_menu = None, None, None, None
user_logo_canvas, active_login_label, sign_out_button, user_canvas = None, None, None, None
user_doc_label, srs_doc_label, srs_image, srs_canvas, srs_canvas_user, match_menu = None, None, None, None, None, None
srs_doc, user_doc, text_view, edit_label, save_button = None, None, None, None, None
ambiguous_menu, compare_heading_menu, anaphoric_canvas, completeness_canvas, excessive_canvas, cordination_canvas= None, None, None, None, None, None
userdoc_view, srsdoc_view, adjective_menu, anaphoric_menu, completeness_menu, excessive_menu, cordination_menu, completeness_label = None, None, None, None, None, None, None, None
anaphoric_label, excessive_label, adj_cordination_label, adj_cordination_canvas  = None, None, None, None
edit_report_word_button, report_user_doc_button, report_label = None, None, None
edit_ambiguous_word_button, ambiguous_user_doc_button, conj_cordination_label = None, None, None
srs_document_path, user_document_path, current_screen, srs_file_name, user_file_name= "", "", "", "", ""
uploaded, is_logged_in = False, False
current_levels = [0] * 10
current_levels2 = [0] * 10
testCase_Document2, SISL, Conj, MWORD, Para, Pronoun, full_srs, full_user, heading_content_list = [], [], [], [], [], [], [], [], []
full_text_Scorce, testCase_Document, Missing, Extra, head_srs, head_user, full_text_Scorce2 = [], [], [], [], [], [], []
tokenized_paragraphs, paragraphs, my_word_tagged, adjective_list, CleanedList, my_word = [], [], [], [], [], []
ant, singleIndividualSentenceList, Noun= [], [], []
singleWords, tagedWords, TQ, testQ, ambi_PC=  [], [], [], [], []
adder, sent= [], []

COLOR_LIST = ["#FF0000", "#00FF00", "#0000FF", "#00FFFF", "#FF00FF", "#C0C0C0", "#800000", "#808000",
              "#008000", "#800080", "#008080", "#000080", "#FF6347", "#006400", "#FFA500", "#FFD700",
              "#00CED1", "#1E90FF", "#8B008B", "#FF1493", "#8B4513", "#DAA520", "#B22222", "#6A5ACD"]

ATOZ = ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u',
        'v', 'w', 'x', 'y', 'z', ".", "," ":", ";", "?", "!", '1', '2', '3', '4', '5', '6', '7', '8', '9', '0']


def destroy_upload_form():
    global srs_document_button, upload_srs_document_label, user_document_button, uploaded_srs_label, uploaded_user_label
    global upload_user_document_label, upload_file_button, srs_canvas, user_doc_label, srs_doc_label, srs_canvas_user

    try:
        srs_document_button.destroy()
        upload_srs_document_label.destroy()
        uploaded_user_label.destroy()
        uploaded_srs_label.destroy()
        user_document_button.destroy()
        upload_user_document_label.destroy()
        upload_file_button.destroy()
        user_doc_label.destroy()
        srs_doc_label.destroy()
        srs_canvas.destroy()
        srs_canvas_user.destroy()
    except AttributeError:
        pass


def destroy_adjective_form():
    global adjective_label, synonym_label, synonym_view, adjective_view
    global srs_canvas, srs_canvas_user

    try:
        adjective_view.destroy()
        adjective_view.frame.destroy()
        adjective_label.destroy()
        synonym_view.frame.destroy()
        synonym_view.destroy()
        if synonym_label is not None:
            synonym_label.destroy()
    except AttributeError as err:
        print(err)

def destroy_cordination():
    global cordination_canvas, adj_cordination_canvas, adj_cordination_label, conj_cordination_label

    try:
        adj_cordination_label.destroy()
        conj_cordination_label.destroy()

        cordination_canvas.frame.destroy()
        cordination_canvas.destroy()
        adj_cordination_canvas.frame.destroy()
        adj_cordination_canvas.destroy()
    except AttributeError as err:
        print("Destroy cordination error:")
        print(err)
               
def cordination():
    global is_logged_in
    global srs_document_path, user_document_path, cordination_canvas, adj_cordination_canvas
    global full_user, conj_cordination_label, adj_cordination_label, NoneType

    if is_logged_in:
        destroy_excessive()
        destroy_completeness()
        destroy_anaphoric()
        destroy_upload_form()
        destroy_login_form()
        destroy_signup_form()
        destroy_match_form()
        destroy_adjective_form()
        destroy_review_form()
        destroy_report_form()
        destroy_ambiguous_form()
        

        if len(srs_document_path) > 0 and len(user_document_path) > 0:
            all_paras = user_doc.paragraphs
            for i in all_paras:
                paragraphs.append(i.text)
            for i in Para:
                MWORD.append(nltk.sent_tokenize(i))
            for i in MWORD:
                for j in i:
                    a = nltk.word_tokenize(j)
                    SISL.append(a)
            for i in SISL:
                a = nltk.pos_tag(i)
                ant.append(a)
            cordination_canvas = ScrolledText(win, width=70, wrap=WORD)
            for i in ant:
                conjunction = 0
                for j in i:
                    if (j[-1] == 'CC'):
                        conjunction=conjunction + 1 
                        Conj.append(j[0])
                    else:
                        continue
                tag = "conj"
                tag = "tomato"
                tag = "black"
                if conjunction>1:
                    cordination_canvas.insert(INSERT, K[0]+' ', 'Conj')
                    cordination_canvas.tag_config('Conj', font=Font(size=13, weight="bold"), foreground= "navy")
                    for K in i:
                        if(K[0] in Conj):
                            cordination_canvas.insert(INSERT, K[0]+' ', 'Tomato')
                            cordination_canvas.tag_config('Tomato', font=Font(size=13, weight="bold"), foreground= "Tomato")
                        else:
                            cordination_canvas.insert(INSERT, K[0]+' ', 'Conj')
                            cordination_canvas.tag_config('Conj', font=Font(size=13, weight="bold"), foreground= "navy")
                else:
                    for K in i:
                        cordination_canvas.insert(INSERT, K[0]+' ', 'Black')
                        cordination_canvas.tag_config('Black', foreground= "Black")
                        
                cordination_canvas.insert(INSERT, '\n')                        
            
            cordination_canvas.pack()
            cordination_canvas.place(x=25, y=251)
            cordination_canvas.config(state=DISABLED)
            conj_cordination_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                                           fg='black', text="Conjunction Cordination")
            conj_cordination_label.pack()
            conj_cordination_label.place(x=210, y=210)
            
            adj_cordination_canvas = ScrolledText(win, width=70, wrap=WORD)
            bb= ' '
            for i in all_paras:
                paragraphs.append(i.text)
            b = ' '.join(paragraphs)
            SentenceS = nltk.sent_tokenize(b)
            for i in SentenceS:
                a = nltk.word_tokenize(i)
                singleWords.append(a)
            for i in singleWords:
                b = nltk.pos_tag(i)
                tagedWords.append(b)
            for i in tagedWords:
                for j in i:
                    if(j[-1] == 'JJ'):
                        testQ.append(i[len(j):len(j)+3])
            for i in testQ:
                if i not in TQ:
                    TQ.append(i)
            for i in TQ:
                for j in i:
                    adder.append(j[-1])
                    sent.append(j[0])
                if (adder[0] == 'NNS' or adder[0] == 'NN') and (adder[1] == 'CC') and (adder[2] == 'NNS' or adder[2] == 'NN'):
                    a = ' '.join(sent)
                    ambi_PC.append(a)
                else:
                    continue
            tag = "conj"
            tag = "black"
            for i in SentenceS:
                for j in ambi_PC:
                    if j in i:
                        bb= i
                        adj_cordination_canvas.insert(INSERT, i+' ', 'Conj')
                        adj_cordination_canvas.tag_config('Conj', font=Font(size=13, weight="bold"), foreground= "navy")
                        ambi_PC.pop(ambi_PC.index(j))
                        break
                    adj_cordination_canvas.insert(INSERT, i+' ', 'Black')
                    adj_cordination_canvas.tag_config('Black', foreground= "Black")
                if len(ambi_PC) == 0:
                    if bb!=i:
                        adj_cordination_canvas.insert(INSERT, i+' ', 'Black')
                        adj_cordination_canvas.tag_config('Black', foreground= "Black")
                        
                adj_cordination_canvas.insert(INSERT, '\n')
                
            adj_cordination_canvas.pack()
            adj_cordination_canvas.place(x=675, y=251)
            adj_cordination_canvas.config(state=DISABLED)
            adj_cordination_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                                          fg='black', text="Adjective Cordination")
            adj_cordination_label.pack()
            adj_cordination_label.place(x=885, y=210)

    else:
        showinfo("Error", "Please log in first")

def destroy_vertical():
    vertical.destroy()
    
def destroy_Horizontal():
    Horizontal.destroy()
    
def destroy_excessive():
    global excessive_canvas, excessive_label

    try:
        excessive_label.destroy()

        excessive_canvas.frame.destroy()
        excessive_canvas.destroy()
    except AttributeError as err:
        print("Destroy anapghoric error:")
        print(err)    
    
def excessive():
    global is_logged_in
    global srs_document_path, user_document_path, excessive_canvas
    global full_user, excessive_label, NoneType

    if is_logged_in:
        destroy_completeness()
        destroy_cordination()
        destroy_upload_form()
        destroy_login_form()
        destroy_signup_form()
        destroy_match_form()
        destroy_adjective_form()
        destroy_review_form()
        destroy_report_form()
        destroy_ambiguous_form()
        destroy_anaphoric()
        

        if len(srs_document_path) > 0 and len(user_document_path) > 0:
            all_paras = user_doc.paragraphs
            for i in all_paras:
                paragraphs.append(i.text)
            for i in paragraphs:
                my_word.append(nltk.sent_tokenize(i))
            excessive_canvas = ScrolledText(win, width=110, wrap=WORD)
            tag = "grey"
            for i in my_word:
                for j in i:
                    a = nltk.word_tokenize(j)
                    singleIndividualSentenceList.append(a)
            for i in singleIndividualSentenceList:
                for j in i:
                    if len(i)<60:
                        excessive_canvas.insert(INSERT, j+' ')
                    else:
                        excessive_canvas.insert(INSERT, j+' ', 'Grey')
                        excessive_canvas.tag_config('Grey', font=Font(size=13, weight="bold"), foreground= "Dark Slate Grey")
                    
                excessive_canvas.insert(INSERT, '\n')
            
            excessive_canvas.pack()
            excessive_canvas.place(x=190, y=240)
            excessive_canvas.config(state=DISABLED)

    else:
        showinfo("Error", "Please log in first")

    excessive_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                            fg='black', text="Excessive Sentences in User Document are:")
    excessive_label.pack()
    excessive_label.place(x=455, y=195)   

    
def destroy_completeness():
    global completeness_canvas, completeness_label

    try:
        completeness_label.destroy()

        completeness_canvas.frame.destroy()
        completeness_canvas.destroy()
    except AttributeError as err:
        print("Destroy anapghoric error:")
        print(err)
        
def completeness():
    global is_logged_in
    global srs_document_path, user_document_path, completeness_canvas
    global full_user, completeness_label, NoneType

    if is_logged_in:
        destroy_excessive()
        destroy_cordination()
        destroy_upload_form()
        destroy_login_form()
        destroy_signup_form()
        destroy_match_form()
        destroy_adjective_form()
        destroy_review_form()
        destroy_report_form()
        destroy_ambiguous_form()
        destroy_anaphoric()
        

        if len(srs_document_path) > 0 and len(user_document_path) > 0:
            all_paras = user_doc.paragraphs
            for i in all_paras:
                paragraphs.append(i.text)
            for i in paragraphs:
                my_word.append(nltk.sent_tokenize(i))
            completeness_canvas = ScrolledText(win, width=110, wrap=WORD)
            tag = "purple"
            for i in my_word:
                for j in i:
                    a = nltk.word_tokenize(j)
                    if (('If' in a) or ('if' in a)) and (('else' not in j) or ('Else' not in j)):
                        completeness_canvas.insert(INSERT, j+' ', 'Purple')
                        completeness_canvas.tag_config('Purple', font=Font(size=13, weight="bold"), foreground= "Purple")
                    else:
                        completeness_canvas.insert(INSERT, j+' ', 'Black')
                        completeness_canvas.tag_config('Black', foreground= "Black")          
            
            
                completeness_canvas.insert(INSERT, '\n')
            
            completeness_canvas.pack()
            completeness_canvas.place(x=190, y=240)
            completeness_canvas.config(state=DISABLED)

    else:
        showinfo("Error", "Please log in first")

    completeness_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                            fg='black', text="Completeness Sentences in User Document are:")
    completeness_label.pack()
    completeness_label.place(x=455, y=195)

        
def destroy_anaphoric():
    global anaphoric_canvas, anaphoric_label

    try:
        anaphoric_label.destroy()

        anaphoric_canvas.frame.destroy()
        anaphoric_canvas.destroy()
    except AttributeError as err:
        print("Destroy anapghoric error:")
        print(err)

def anaphoric():
    global is_logged_in
    global srs_document_path, user_document_path, anaphoric_canvas
    global full_user, anaphoric_label, NoneType

    if is_logged_in:
        destroy_cordination()
        destroy_upload_form()
        destroy_login_form()
        destroy_signup_form()
        destroy_match_form()
        destroy_adjective_form()
        destroy_review_form()
        destroy_report_form()
        destroy_ambiguous_form()
        destroy_completeness()
        

        if len(srs_document_path) > 0 and len(user_document_path) > 0:
            all_paras = user_doc.paragraphs
            for i in all_paras:
                paragraphs.append(i.text)
            for i in paragraphs:
                my_word.append(nltk.sent_tokenize(i))
            for i in my_word:
                for j in i:
                    a = nltk.word_tokenize(j)
                    singleIndividualSentenceList.append(a)
            for i in singleIndividualSentenceList:
                a = nltk.pos_tag(i)
                ant.append(a)
            anaphoric_canvas = ScrolledText(win, width=110, wrap=WORD)
            for i in ant:
                noun = 0
                pronoun = 0
                
                for j in i:
                    if (j[-1] == 'NN') | (j[-1] == 'NNS') | (j[-1] == 'NNP') | (j[-1] == 'NNPS'):
                        noun=noun + 1
                        Noun.append(j[0])
                    elif (j[-1] == 'PRP') | (j[-1] == 'PRP$') | (j[-1] == 'WP') | (j[-1] == 'WP$'):
                        if(len(Noun)>1):
                            pronoun=pronoun + 1
                            Pronoun.append(j[0])
                    else:
                        continue
                            
                tag = "red"
                tag = "grey"
                tag = "blue"
                if noun>1 & pronoun>=1:
                    for l in i:
                        if(l[0] in Noun):
                            anaphoric_canvas.insert(INSERT, l[0]+' ', 'Red')
                            anaphoric_canvas.tag_config('Red', font=Font(size=13, weight="bold"), foreground= "Red")
                        elif(l[0] in Pronoun):
                            anaphoric_canvas.insert(INSERT, l[0]+' ', 'Grey')
                            anaphoric_canvas.tag_config('Grey', font=Font(size=13, weight="bold"), foreground= "Grey")
                            Noun.clear()
                        else:
                            anaphoric_canvas.insert(INSERT, l[0]+' ', 'Blue')
                            anaphoric_canvas.tag_config('Blue', font=Font(size=13, weight="bold"), foreground= "Blue2")
                else:
                    for l in i:
                        anaphoric_canvas.insert(INSERT, l[0]+' ', 'Black')
                        anaphoric_canvas.tag_config('Black', foreground= "Black")
                    
                    anaphoric_canvas.insert(INSERT, '\n')
            
            anaphoric_canvas.pack()
            anaphoric_canvas.place(x=190, y=240)
            anaphoric_canvas.config(state=DISABLED)

    else:
        showinfo("Error", "Please log in first")

    anaphoric_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                            fg='black', text="Anaphoric Sentences in User Document are:")
    anaphoric_label.pack()
    anaphoric_label.place(x=455, y=195)


def adjective_synonyms():
    global adjective_label,synonym_label,synonym_view,adjective_view
    global is_logged_in, current_screen, user_doc_label, srs_doc_label

    if current_screen != "Adjective_synonyms":
        current_screen = "Adjective_synonyms"
        destroy_upload_form()
        destroy_login_form()
        destroy_review_form()
        destroy_signup_form()
        destroy_match_form()
        destroy_compare_headings_form()
        destroy_ambiguous_form()
        destroy_report_form()
        destroy_excessive()
        destroy_completeness()
        destroy_cordination()
        destroy_anaphoric()

        all_paras = user_doc.paragraphs
        regex = r'[>“"<"'']'
        for i in all_paras:                                 # get text from all paragraphs
            paragraphs.append(i.text)
        for i in range(len(paragraphs)):                    #convert text into lower
            paragraphs[i] = paragraphs[i].lower()
        for i in paragraphs:                                 #apply word tokenizer of all
            tokenized_paragraphs.append(nltk.word_tokenize(i))
        for i in range(len(paragraphs)):                                             #stop words removal
            words = nltk.word_tokenize(paragraphs[i])                                   # Tokenize into words.
            words = [x for x in words if x not in stopwords.words('english')]         # Remove the stop words.
            paragraphs[i] = ' '.join(words)
        for i in tokenized_paragraphs:                                #apply pos Tag on the tokenized bag of words
            my_word_tagged.append(nltk.pos_tag(i))
        for i in my_word_tagged:                                       #find the adjectives
            current = i
            for j in current:
                if 'JJ' in j or 'JJR' in j or 'JJS' in j:
                    if j[0] not in adjective_list:
                        adjective_list.append(j[0])
                        
        adjective_view = ScrolledText(win, width=70, wrap=WORD)

        for i in adjective_list:                                          #remove the sybmols
            CleanedList.append(re.sub(regex,'',i))
        for i in CleanedList:
            if i == '':
                CleanedList.remove(i)
            elif len(i) == 1:
                CleanedList.remove(i)
            else:
                continue

        for i in CleanedList:                      
            adjective_view.insert(INSERT, i+'\n')
            
        adjective_view.pack()
        adjective_view.place(x=25, y=251)
        adjective_view.config(state=DISABLED)

        adjective_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                                   fg='black', text="Adjectives")
        adjective_label.pack()
        adjective_label.place(x=240, y=210)

        synonym_view = ScrolledText(win, width=70, wrap=WORD)

        synonyms = []
        for adjectives in CleanedList:
            for syn in wordnet.synsets(adjectives):
                for lm in syn.lemmas():
                     synonyms.append(lm.name())
            a = list(set(synonyms))
            synonym_view.tag_config('Red', font=Font(size=11, weight="bold"), foreground='medium violet red')
            synonym_view.insert(INSERT,adjectives+': \n','Red')
            synonym_view.insert(INSERT, str(list(set(synonyms)))+'\n')
            synonym_view.insert(INSERT, 'This '+adjectives+' having '+str(len(a))+' different meaning and system has detected it as ambiguous word\n\n',"Blue")
            synonym_view.tag_config('Blue', font=Font(size=11, weight="bold"), foreground='Blue1')

        synonym_view.pack()
        synonym_view.place(x=675, y=251)
        synonym_view.config(state=DISABLED)

        synonym_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                                   fg='black', text="Synonyms")
        synonym_label.pack()
        synonym_label.place(x=915, y=210)


def destroy_login_form():
    global username_label, username_entry, password_label, password_entry, login_button

    try:
        login_button.destroy()
        username_label.destroy()
        username_entry.destroy()
        password_label.destroy()
        password_entry.destroy()

        win.unbind_all("<Return>")
    except AttributeError:
        pass


def destroy_review_form():
    global srs_canvas, user_canvas, user_doc_label, srs_doc_label
    global srs_document_path, user_document_path

    try:
        user_canvas.frame.destroy()
        srs_canvas.frame.destroy()

        user_canvas.destroy()
        srs_canvas.destroy()
        user_doc_label.destroy()
        srs_doc_label.destroy()

        # srs_document_path, user_document_path = "", ""

    except AttributeError:
        pass


def destroy_signup_form():
    global username_label, username_entry, password_label, password_entry, sign_up_button
    global password_entry_2, password_label_2

    try:
        username_label.destroy()
        sign_up_button.destroy()
        username_entry.destroy()
        password_label.destroy()
        password_entry.destroy()
        password_entry_2.destroy()
        password_label_2.destroy()
    except AttributeError:
        pass


def destroy_match_form():
    global text_view, userdoc_view, srsdoc_view, srs_canvas, srs_canvas_user, user_doc_label, srs_doc_label

    try:
        srs_canvas.destroy()
        srs_canvas_user.destroy()

        user_doc_label.destroy()
        srs_doc_label.destroy()
        userdoc_view.destroy()
        userdoc_view.frame.destroy()
        srsdoc_view.destroy()
        srsdoc_view.frame.destroy()
        text_view.frame.destroy()
    except Exception as err:
        pass


def destroy_ambiguous_form():
    global text_view, edit_ambiguous_word_button, ambiguous_user_doc_button, edit_label, save_button, ambiguous_label
    try:
        ambiguous_label.destroy()
        edit_ambiguous_word_button.destroy()
        ambiguous_user_doc_button.destroy()

        text_view.frame.destroy()
        text_view.destroy()

        edit_label.destroy()
        save_button.destroy()
    except AttributeError:
        pass


def destroy_compare_headings_form():
    global text_view, srs_canvas, srs_canvas_user
    try:

        text_view.frame.destroy()
        text_view.destroy()
        srs_canvas.destroy()
        srs_canvas_user.destroy()

    except AttributeError:
        pass


def destroy_report_form():
    global edit_report_word_button, report_user_doc_button, text_view, edit_label, report_label

    try:
        report_label.destroy()
        text_view.destroy()
        edit_label.destroy()
        edit_report_word_button.destroy()
        report_user_doc_button.destroy()
    except Exception as err:
        print("Destroy report form error:")
        print(err)
        
def sign_out():
    global user_logo_canvas, active_login_label, sign_out_button, upload_menu, srs_canvas, srs_canvas_user
    global or_label, is_logged_in, login_menu, sign_up_menu_button, current_screen
    global review_menu, match_menu, ambiguous_menu, compare_heading_menu, vertical, horizontal
    global srs_document_path, user_document_path, uploaded, adjective_menu, anaphoric_menu, excessive_menu, completeness_menu, cordination_menu

    uploaded = False
    showinfo("Info", "Signed Out Successfully")

    try:
        # asd
        upload_menu.destroy()
        sign_out_button.destroy()
        user_logo_canvas.destroy()
        active_login_label.destroy()
        cordination_menu.destroy()
        excessive_menu.destroy()
        completeness_menu.destroy()
        anaphoric_menu.destroy()
        adjective_menu.destroy()
        user_logo_canvas.destroy()
        active_login_label.destroy()
        upload_menu.destroy()
        review_menu.destroy()
        match_menu.destroy()
        ambiguous_menu.destroy()
        destroy_compare_headings_form()
        destroy_adjective_form()
        sign_out_button.destroy()
        compare_heading_menu.destroy()
        srs_canvas.destroy()
        srs_canvas_user.destroy()
        destroy_ambiguous_form()
        destroy_completeness()
        destroy_anaphoric()
        destroy_cordination()
        destroy_excessive()
        destroy_vertical()
        destroy_horizontal()
    except AttributeError:
        pass
    finally:
        srs_document_path, user_document_path = "", ""

    login_menu = Button(win, font=("Californian FB", 11, 'bold '), bg='SteelBlue2', fg='Midnight Blue', width=14,
                        text="Login",
                        activebackground='SlateBlue2',
                        activeforeground='Midnight Blue',
                        command=lambda:[show_login_form(), clickme2()])
    login_menu.pack()
    login_menu.place(x=433, y=85)

    sign_up_menu_button = Button(win, font=("Californian FB", 12, 'bold '), bg='SteelBlue2', fg='Midnight Blue',
                                 width=14,
                                 text="Sign Up", activebackground='SlateBlue2', activeforeground='Midnight Blue',
                                 command=lambda:[show_signup_form(), clickme3()])
    sign_up_menu_button.pack()
    sign_up_menu_button.place(x=713, y=85)

    show_login_form()

    or_label = Label(win, font=("Californian FB", 16, 'bold '), bg='Light Sky Blue', fg='Midnight Blue', text="OR")
    or_label.pack()
    or_label.place(x=627, y=85)

    is_logged_in = False
    current_screen = ""


def show_user_icon(uname: str):
    global user_logo_canvas, active_login_label, sign_out_button

    user_logo_canvas = Canvas(height=110, width=110, highlightthickness=0, bg='Light Sky Blue')

    user_logo_canvas.create_image(50, 50, image=logo_image)
    user_logo_canvas.pack()
    user_logo_canvas.place(x=1135, y=36)

    active_login_label = Label(win, font=("Californian FB", 14, 'bold '), bg='Light Sky Blue', fg='Midnight Blue',
                               text=f"Greetings, {uname}")
    active_login_label.pack()
    active_login_label.place(x=1110, y=125)

    sign_out_button = Button(win, font=("Californian FB", 12, 'bold '), bg='RoyalBlue4', fg='black', width=12,
                             activebackground='SlateBlue2', activeforeground='Midnight Blue',
                             text="Sign Out", command=sign_out)
    sign_out_button.pack()
    sign_out_button.place(x=575, y=675)
    

def show_login_form():
    global or_label, srs_canvas, srs_canvas_user, username_label, username_entry, password_label, password_entry
    global login_button, upload_menu, current_screen

    def validate_login(uname, pwd):
        global is_logged_in, user_logo_canvas, active_login_label, sign_out_button, upload_menu

        user_list = pandas.read_csv("user.csv").to_dict(orient="records")

        for user in user_list:
            user_flag = str(user["username"]).lower() == str(uname.get()).lower()
            pwd_flag = str(user["password"]) == str(pwd.get())
            flag = user_flag and pwd_flag

            if flag:
                is_logged_in = True
                break
            else:
                is_logged_in = False

        if is_logged_in:
            showinfo("Success", f"Welcome {uname.get()}")
            show_user_icon(uname.get())

            uname.set("")
            pwd.set("")

            destroy_login_form()
            login_menu.destroy()
            or_label.destroy()
            sign_up_menu_button.destroy()
            upload_menu = Button(win, font=("Californian FB", 11, 'bold '), bg='SteelBlue2', fg='Midnight Blue',
                                 width=12,
                                 text="Upload", activebackground='SlateBlue2', activeforeground='Midnight Blue',
                                 command=lambda:[show_upload_form(),clickme4()])
            upload_menu.pack()
            upload_menu.place(x=575, y=75)

        else:
            showinfo("Error", f"Username or password didn't match")

    if current_screen != "login":
        current_screen = "login"

        destroy_upload_form()
        destroy_review_form()
        destroy_signup_form()
        destroy_match_form()

        username = StringVar()
        password = StringVar()

        username_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                               fg='Midnight Blue',
                               text="Enter User Name")
        username_label.pack()
        username_label.place(x=578, y=160)

        username_entry = Entry(win, textvariable=username)
        username_entry.pack()
        username_entry.place(x=585, y=207)
        username_entry.focus()

        password_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                               fg='Midnight Blue',
                               text="Enter Password")
        password_label.pack()
        password_label.place(x=583, y=247)

        password_entry = Entry(win, textvariable=password, show='*')
        password_entry.pack()
        password_entry.place(x=585, y=295)

        validate_login = partial(validate_login, username, password)

        login_button = Button(win, font=("Californian FB", 12, 'bold '), bg='SteelBlue2', fg='Midnight Blue',
                              text="   Login   ", command=validate_login, activebackground='SlateBlue2',
                              activeforeground='Midnight Blue')
        login_button.pack()
        login_button.place(x=610, y=357)


def open_file():
    file_path = askopenfile(mode='r', filetypes=[('Word Document Files', '*docx', '')])
    if file_path is not None:
        pass


def open_user_file():
    global user_document_path, user_file_name

    file_path = askopenfile(mode='r', filetypes=[('Word Document Files', '*docx')])
    if file_path is not None:
        user_document_path = file_path.name
        user_file_name = user_document_path.split("/")[-1]


def open_srs():
    global srs_document_path, srs_file_name

    file_path = askopenfile(mode='r', filetypes=[('Word Document Files', '*docx')])
    if file_path is not None:
        srs_document_path = file_path.name
        srs_file_name = srs_document_path.split("/")[-1]


def get_ambiguous_word_user_doc():
    global user_doc, edit_label, save_button

    if edit_label is not None:
        edit_label.destroy()

    if save_button is not None:
        save_button.destroy()

    ambiguous_words = read_ambiguous_word()

    if len(ambiguous_words) > 0:
        match_words = ["Ambiguous Words in User Doc:\n"]

        for p in user_doc.paragraphs:
            para = p.text.split(" ")
            for w in para:
                if w.lower() in ambiguous_words:
                    match_words.append(w.lower())

        match_words_fin = list(dict.fromkeys(match_words))
        write_text_view(match_words_fin)
    else:
        messagebox.showerror("Error", "Ambiguous Word DB is empty")


def show_ambiguous_word():
    texts = ["Ambiguous Words in XLSX File:\n"]

    ambiguous_words = read_ambiguous_word()
    for w in ambiguous_words:
        texts.append(w)

    write_text_view(texts)


def write_text_view(texts: list):
    global text_view

    if text_view is not None:
        text_view.frame.destroy()
        text_view.destroy()

    text_view = ScrolledText(win, width=110, wrap=WORD)

    first = True
    for text in texts:
        if first:
            line = f"{text}"
            first = False
        else:
            line = f"- {text}\n"
        text_view.insert(INSERT, line)

    text_view.pack()
    text_view.place(x=190, y=240)
    text_view.config(state=DISABLED)


def read_ambiguous_word() -> list:
    wb = xl.load_workbook("Ambiguous_DB.xlsx")
    ws = wb.worksheets[0]

    i = 2
    ambiguous_words = []
    word = "word"
    while word is not None:
        word = ws.cell(i, 1).value
        if word is not None:
            ambiguous_words.append(word.lower())
        i += 1

    return ambiguous_words


def save_ambiguous_word():
    global text_view, edit_label, save_button

    confirm_save = messagebox.askyesno("Confirmation", "Save ambiguous words?")
    if confirm_save:
        word_list = []
        text_val = text_view.get("1.0", END).split("\n")
        temp_val = [x for x in text_val if len(x) > 0 and 'XLSX' not in x]
        for x in temp_val:
            z = x.split(" ")[1]
            word_list.append(z.lower())

        distinct_list = np.unique(word_list)

        wb = xl.load_workbook("Ambiguous_DB.xlsx")
        ws = wb.worksheets[0]

        ws.delete_rows(2, 11000)

        i = 2
        for word in distinct_list:
            ws.cell(i, 1).value = word
            i += 1

        messagebox.showinfo("Info", "Successfully updated ambiguous word list")
        wb.save("Ambiguous_DB.xlsx")
        show_ambiguous_word()

        edit_label.destroy()
        save_button.destroy()


def edit_ambiguous_word():
    global edit_label, text_view, save_button

    show_ambiguous_word()

    edit_label = Label(win, text="Edit Mode", font=Font(size=14, weight="bold"), bg='Light Sky Blue',
                       fg='Midnight Blue')
    edit_label.pack()
    edit_label.place(x=191, y=195)

    text_view.config(state=NORMAL)

    save_button = Button(win, font=("Californian FB", 11, 'bold '), bg='SteelBlue2', fg='Midnight Blue', width=18,
                         activebackground='SlateBlue2', activeforeground='Midnight Blue',
                         text="Save",
                         command=lambda: save_ambiguous_word())
    save_button.pack()
    save_button.place(x=920, y=650)


def show_ambiguous_form():
    global edit_ambiguous_word_button, ambiguous_user_doc_button, text_view, edit_label, current_screen, ambiguous_label

    if current_screen != "ambiguous":
        current_screen = "ambiguous"

        destroy_upload_form()
        destroy_login_form()
        destroy_review_form()
        destroy_signup_form()
        destroy_match_form()
        destroy_adjective_form()
        destroy_completeness()
        destroy_compare_headings_form()
        destroy_anaphoric()
        destroy_cordination()
        destroy_excessive()
        destroy_report_form()
        
        
        ambiguous_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                            fg='black', text="Ambiguous words in User Document are:")
        ambiguous_label.pack()
        ambiguous_label.place(x=475, y=195)
        edit_ambiguous_word_button = Button(win, font=("Californian FB", 11, 'bold '), bg='SteelBlue2',
                                            fg='Midnight Blue',
                                            activebackground='SlateBlue2', activeforeground='Midnight Blue',
                                            text=" Edit Ambiguous List ",
                                            command=lambda:[edit_ambiguous_word(), clickme12()])

        ambiguous_user_doc_button = Button(win, font=("Californian FB", 11, 'bold '), bg='SteelBlue2',
                                           fg='Midnight Blue',
                                           activebackground='SlateBlue2', activeforeground='Midnight Blue',
                                           text=" Amb Word in User Doc ",
                                           command=lambda:[get_ambiguous_word_user_doc(), clickme13()])

        text_view = ScrolledText(win, width=110, wrap=WORD)
        text_view.pack()
        text_view.place(x=190, y=240)

        edit_ambiguous_word_button.pack()
        ambiguous_user_doc_button.pack()

        edit_ambiguous_word_button.place(x=20, y=270)
        ambiguous_user_doc_button.place(x=10, y=370)

        show_ambiguous_word()
        
        

def extract_heading(level: str):
    global user_doc, heading_content_list, text_view

    style = f"Heading {level}"
    heading_content_list = []
    for p in user_doc.paragraphs:
        if p.style.name == style:
            heading_content_list.append(p.text)
            heading_content_list.append("\n")
            heading_content_list.append("\n")

    if text_view is not None:
        text_view.frame.destroy()
        text_view.destroy()
    text_view = ScrolledText(win, width=110, wrap=WORD)
    for text in heading_content_list:
        text_view.insert(INSERT, text)

    text_view.pack()
    text_view.place(x=190, y=240)
    text_view.config(state=DISABLED)


def match_heading_user_srs():
    global srs_doc, user_doc, userdoc_view, srsdoc_view
    global is_logged_in, current_screen, user_doc_label, srs_doc_label
    global srs_canvas, srs_canvas_user

    srs_canvas.destroy()
    srs_canvas_user.destroy()

    if is_logged_in:
        if current_screen != "match":
            current_screen = "match"

            destroy_upload_form()
            destroy_completeness()
            destroy_excessive()
            destroy_anaphoric()
            destroy_cordination()
            destroy_login_form()
            destroy_signup_form()
            destroy_review_form()
            destroy_ambiguous_form()
            destroy_compare_headings_form()
            destroy_adjective_form()
            destroy_report_form()

            temp_srs_heading_list = []
            for para in testCase_Document:
                temp_srs_heading_list.append(para)
            srs_heading_list = list(dict.fromkeys(temp_srs_heading_list))

            srsdoc_view = ScrolledText(win, width=70, wrap=WORD)
            userdoc_view = ScrolledText(win, width=70, wrap=WORD)

            user_headings = []
            for para in testCase_Document2:
                user_headings.append(para.replace("\n", ""))

            srs_headings = []
            for para in testCase_Document:
                if para in srs_heading_list:
                    srs_headings.append(para.replace("\n", ""))

            similar_headings = []
            for uh in user_headings:
                if uh in srs_headings:
                    similar_headings.append(uh)

            r = 0
            userdoc_view.delete('1.0', END)
            for h in user_headings:
                if h in similar_headings:
                    userdoc_view.insert(INSERT, h + '\n\n', "Midnight Blue")
                    userdoc_view.tag_config('Midnight Blue', foreground='Midnight Blue')
                else:
                    userdoc_view.insert(INSERT, h + '\n\n', "Red")
                    userdoc_view.tag_config('Red', font=Font(size=11, weight="bold"), foreground='Red')
                    Missing.append(h)

                r = r + 1

            srsdoc_view.delete('1.0', END)
            for h in srs_headings:
                if h in similar_headings:
                    srsdoc_view.insert(INSERT, h + '\n\n', "Midnight Blue")
                    srsdoc_view.tag_config('Midnight Blue', foreground='Midnight Blue')
                else:
                    srsdoc_view.insert(INSERT, h + '\n\n', "Red")
                    srsdoc_view.tag_config('Red', font=Font(size=11, weight="bold"), foreground='Red')
                    Extra.append(h)
                r = r + 1

            userdoc_view.pack()
            userdoc_view.place(x=25, y=251)
            userdoc_view.config(state=DISABLED)

            srsdoc_view.pack()
            srsdoc_view.place(x=675, y=251)
            srsdoc_view.config(state=DISABLED)

            user_doc_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                                   fg='black', text="User Document")
            user_doc_label.pack()
            user_doc_label.place(x=240, y=210)

            srs_doc_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                                  fg='black', text="SRS Document")
            srs_doc_label.pack()
            srs_doc_label.place(x=915, y=210)
    

def show_upload_form():
    global srs_document_button, upload_srs_document_label, user_document_button, review_menu
    global uploaded_srs_label, uploaded_user_label
    global upload_user_document_label, upload_file_button, srs_canvas, srs_canvas_user, srs_image
    global user_doc_label, srs_doc_label
    global user_doc, compare_heading_menu, current_screen

    if is_logged_in:
        if current_screen != "upload":
            current_screen = "upload"

            destroy_login_form()
            destroy_cordination()
            destroy_completeness()
            destroy_excessive()
            destroy_anaphoric()
            destroy_review_form()
            destroy_signup_form()
            destroy_match_form()
            destroy_ambiguous_form()
            destroy_adjective_form()

            def upload_file():
                global srs_document_path, srs_image, srs_canvas, user_document_path, review_menu, srs_canvas_user
                global full_srs, full_user, srs_canvas_user, srs_canvas, user_canvas, srs_doc, user_doc, match_menu
                global user_doc_label, srs_doc_label, uploaded
                global ambiguous_menu, compare_heading_menu, current_levels, current_levels2
                global full_text_Scorce, full_text_Scorce2, testCase_Document, testCase_Document2, adjective_menu, excessive_menu, completeness_menu, anaphoric_menu, cordination_menu

                lab = Label(win, font=("Californian FB", 14, 'bold '), fg='Midnight Blue',
                            text='File Uploaded Successfully!',
                            bg='Light Sky Blue')
                lab.place(x=527, y=505)

                def remove_label():
                    time.sleep(1)
                    lab.destroy()

                th = Thread(target=remove_label)
                th.start()

                copy(srs_document_path, "./files/")
                copy(user_document_path, "./files/")

                srs_doc = Document(srs_document_path)

                def format_levels(cur_lev):
                    levs = [str(cl) for cl in cur_lev if cl != 0]
                    return '.'.join(levs)  # Customize your format here

                full_text_Scorce = []
                if len(testCase_Document) > 0:
                    testCase_Document = []
                    current_levels = [0] * 10
                if len(testCase_Document2) > 0:
                    testCase_Document2 = []
                    current_levels2 = [0] * 10

                for p in srs_doc.paragraphs:
                    if p.style.name not in level_from_style_name:
                        full_text_Scorce.append(p.text)
                    else:
                        level = level_from_style_name[p.style.name]
                        current_levels[level] += 1
                        for lev in range(level + 1, 10):
                            current_levels[lev] = 0
                        testCase_Document.append(format_levels(current_levels) + ' ' + p.text)

                full_srs = []
                for x in srs_doc.paragraphs:
                    full_srs.append(x.text)
                    full_srs.append("\n")
                    full_srs.append("\n")

                user_doc = Document(user_document_path)

                full_text_Scorce2 = []
                for p in user_doc.paragraphs:
                    if p.style.name not in level_from_style_name2:
                        full_text_Scorce2.append(p.text)
                    else:
                        level = level_from_style_name2[p.style.name]
                        current_levels2[level] += 1
                        for lev in range(level + 1, 10):
                            current_levels2[lev] = 0
                        testCase_Document2.append(format_levels(current_levels2) + ' ' + p.text)

                full_user = []
                for x in user_doc.paragraphs:
                    full_user.append(x.text)
                    full_user.append("\n")
                    full_user.append("\n")

                if not uploaded:
                    review_menu = Button(win, font=("Californian FB", 11, 'bold '), bg='SteelBlue2', fg='Midnight Blue',
                                         width=12,
                                         text="Review", activebackground='SlateBlue2', activeforeground='Midnight Blue',
                                         command=lambda:[show_review_form(), clickme5()])
                    review_menu.pack()
                    review_menu.place(x=275, y=75)

                    match_menu = Button(win, font=("Californian FB", 11, 'bold '), bg='SteelBlue2', fg='Midnight Blue',
                                        width=12,
                                        text="Match Heading", activebackground='SlateBlue2',
                                        activeforeground='Midnight Blue',
                                        command=lambda:[match_heading_user_srs(), clickme6()])
                    match_menu.pack()
                    match_menu.place(x=425, y=75)

                    ambiguous_menu = Button(win, font=("Californian FB", 11, 'bold '), bg='SteelBlue2',
                                            fg='Midnight Blue',
                                            width=12,
                                            text="Ambiguous", activebackground='SlateBlue2',
                                            activeforeground='Midnight Blue',
                                            command=lambda:[show_ambiguous_form(), clickme9()])
                    ambiguous_menu.pack()
                    ambiguous_menu.place(x=875, y=75)

                    compare_heading_menu = Button(win, font=("Californian FB", 11, 'bold '), bg='SteelBlue2',
                                                  fg='Midnight Blue',
                                                  width=12,
                                                  text="Report", activebackground='SlateBlue2',
                                                  activeforeground='Midnight Blue', command=lambda:[show_report_form(), clickme8()])
                    compare_heading_menu.pack()
                    compare_heading_menu.place(x=725, y=75)

                    adjective_menu = Button(win, font=("Californian FB", 11, 'bold '), bg='SteelBlue2',
                                            fg='Midnight Blue',
                                            width=12,
                                            text="Adjec & Synon", activebackground='SlateBlue2',
                                            activeforeground='Midnight Blue', command=lambda:[adjective_synonyms(), clickme7()])
                    adjective_menu.pack()
                    adjective_menu.place(x=425, y=120)
                    
                    cordination_menu = Button(win, font=("Californian FB", 11, 'bold'),bg='SteelBlue2',
                                            fg='Midnight Blue',
                                            width=12,
                                            text="Cordination", activebackground='SlateBlue2',
                                            activeforeground='Midnight Blue', command=lambda:[cordination(), clickme11()])
                    cordination_menu.pack()
                    cordination_menu.place(x=725, y=120)

                    anaphoric_menu = Button(win, font=("Californian FB", 11, 'bold '), bg='SteelBlue2',
                                            fg='Midnight Blue',
                                            width=12,
                                            text="Anaphoric", activebackground='SlateBlue2',
                                            activeforeground='Midnight Blue', command=lambda:[anaphoric(), clickme10()])
                    anaphoric_menu.pack()
                    anaphoric_menu.place(x=575, y=120)
                    
                    completeness_menu = Button(win, font=("Californian FB", 11, 'bold '), bg='SteelBlue2',
                                            fg='Midnight Blue',
                                            width=12,
                                            text="Completeness", activebackground='SlateBlue2',
                                            activeforeground='Midnight Blue', command=lambda:[clickme15(), completeness()])
                    completeness_menu.pack()
                    completeness_menu.place(x=875, y=120)
                    
                    excessive_menu = Button(win, font=("Californian FB", 11, 'bold '), bg='SteelBlue2',
                                            fg='Midnight Blue',
                                            width=12,
                                            text="Excessive", activebackground='SlateBlue2',
                                            activeforeground='Midnight Blue', command=lambda:[clickme14(), excessive()])
                    excessive_menu.pack()
                    excessive_menu.place(x=275, y=120)
                    
                    vertical =Frame(win, bg='Midnight Blue', height=90,width=5)
                    vertical.place(x=255, y=62)
                    vertical =Frame(win, bg='Midnight Blue', height=91,width=5)
                    vertical.place(x=1010, y=75)
                    horizontal =Frame(win, bg='Midnight Blue', height=5,width=288)
                    horizontal.place(x=255, y=58)
                    horizontal =Frame(win, bg='Midnight Blue', height=5,width=290)
                    horizontal.place(x=725, y=164)

                    uploaded = True
                print(f"uploaded = {uploaded}")

                if srs_canvas is not None:
                    srs_canvas.destroy()
                srs_canvas = Canvas(height=300, width=350, highlightthickness=0, bg="Light Sky Blue")

                if srs_canvas_user is not None:
                    srs_canvas_user.destroy()
                srs_canvas_user = Canvas(height=300, width=350, highlightthickness=0, bg="Light Sky Blue")

                if srs_document_path != "":
                    srs_image = ImageTk.PhotoImage(file="doc.png")
                    srs_canvas.create_image(75, 75, image=srs_image)
                    srs_canvas.create_text(75, 150, text=srs_file_name)

                if user_document_path != "":
                    srs_canvas_user.create_image(75, 75, image=srs_image)
                    srs_canvas_user.create_text(75, 150, text=user_file_name)

                srs_canvas.pack()
                srs_canvas.place(x=935, y=250)
                srs_canvas_user.pack()
                srs_canvas_user.place(x=180, y=250)

            if user_document_path != "" and srs_document_path != "":
                srs_canvas = Canvas(height=300, width=400, highlightthickness=0, bg="Light Sky Blue")
                srs_canvas_user = Canvas(height=300, width=400, highlightthickness=0, bg="Light Sky Blue")
                if srs_document_path != "":
                    srs_image = ImageTk.PhotoImage(file="doc.png")
                    srs_canvas.create_image(75, 75, image=srs_image)
                    srs_canvas.create_text(75, 150, text=srs_file_name)

                if user_document_path != "":
                    srs_canvas_user.create_image(75, 75, image=srs_image)
                    srs_canvas_user.create_text(75, 150, text=user_file_name)

                srs_canvas.pack()
                srs_canvas.place(x=935, y=250)
                srs_canvas_user.pack()
                srs_canvas_user.place(x=180, y=250)

            srs_document_button = Button(win, font=("Californian FB", 11, 'bold '), width=12, bg='SteelBlue2',
                                         fg='Midnight Blue', text='Upload SRS Doc', activebackground='SlateBlue2',
                                         activeforeground='Midnight Blue', command=lambda: open_srs())
            srs_document_button.place(x=575, y=275)

            uploaded_srs_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                                       fg='black', text='Uploaded SRS Document is:')
            uploaded_srs_label.pack()
            uploaded_srs_label.place(x=900, y=210)

            uploaded_user_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                                        fg='black', text='Uploaded User Document is:')
            uploaded_user_label.pack()
            uploaded_user_label.place(x=145, y=210)

            upload_srs_document_label = Label(win, font=("Californian FB", 14, 'bold '), bg='Light Sky Blue',
                                              fg='Midnight Blue', text='Upload IEEE SRS Document')
            upload_srs_document_label.pack()
            upload_srs_document_label.place(x=515, y=210)

            user_document_button = Button(win, font=("Californian FB", 11, 'bold '), width=12, bg='SteelBlue2',
                                          fg='Midnight Blue', text='Upload User Doc', activebackground='SlateBlue2',
                                          activeforeground='Midnight Blue', command=lambda: open_user_file())
            user_document_button.pack()
            user_document_button.place(x=575, y=415)

            upload_user_document_label = Label(win, font=("Californian FB", 14, 'bold '), bg='Light Sky Blue',
                                               fg='Midnight Blue',
                                               text='Upload User Document')
            upload_user_document_label.place(x=540, y=350)

            upload_file_button = Button(win, font=("Californian FB", 11, 'bold '), width=12, bg='dark slate grey', fg='black',
                                        text='Upload Files', activebackground='slate grey', activeforeground='black',
                                        command=upload_file)
            upload_file_button.place(x=575, y=585)
    else:
        showinfo("Error", "Please log in first")
        user_doc_label.destroy()
        srs_doc_label.destroy()
        
        
def show_report_form():
    global edit_report_word_button, report_user_doc_button, text_view, edit_label, current_screen, report_label

    if current_screen != "report":
        current_screen = "report"

        destroy_anaphoric()
        destroy_excessive()
        destroy_completeness()
        destroy_cordination()
        destroy_upload_form()
        destroy_login_form()
        destroy_review_form()
        destroy_signup_form()
        destroy_match_form()
        destroy_compare_headings_form()
        destroy_ambiguous_form()
        destroy_adjective_form()

        tbl = str.maketrans('', '', digits)
        full_report = ['\t\t\tIntelligent Analysis of Software Requirements Specification\t\t\t\n',
                       '\nIEEE Document as Standard : ' + srs_file_name, '\nUser Document : ' + user_file_name]
        j = 1
        full_report.append('\n\nExtra Headings in the user document are as below :\n')
        for i in Missing:
            full_report.append('\n' + str(j) + '. ' + str(i.translate(tbl).replace('.', '').lstrip()))
            j += 1
        full_report.append('\n\nMissing Headings in the user document are as below :\n')
        a = 1
        for k in Extra:
            full_report.append('\n' + str(a) + '. ' + str(k.translate(tbl).replace('.', '').lstrip()))
            a += 1
        now = datetime.datetime.now()
        full_report.append("\n\nCurrent date and time : ")
        full_report.append('\n' + str(now.strftime("%Y-%m-%d")))
        full_report.append('\n' + str(now.strftime("%H:%M:%S")))

        text_view = ScrolledText(win, width=110, wrap=WORD)
        for x in full_report:
            text_view.insert(INSERT, x)

        text_view.pack()
        text_view.place(x=190, y=240)

        report_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                             fg='black', text="User Document Report is:")
        report_label.pack()
        report_label.place(x=542, y=195)


def show_review_form():
    global open_button, is_logged_in, srs_canvas, srs_image, user_doc_label, srs_doc_label
    global srs_document_path, user_document_path, user_canvas
    global full_srs, full_user, srs_doc, user_doc, current_screen, adjective_list

    if is_logged_in:
        if current_screen != "review":
            current_screen = "review"

            destroy_upload_form()
            destroy_cordination()
            destroy_completeness()
            destroy_excessive()
            destroy_anaphoric()
            destroy_login_form()
            destroy_signup_form()
            destroy_match_form()
            destroy_ambiguous_form()
            destroy_compare_headings_form()
            destroy_adjective_form()
            destroy_report_form()

            if len(srs_document_path) > 0 and len(user_document_path) > 0:
                ambiguous_list = read_ambiguous_word()
                user_canvas = ScrolledText(win, width=70, wrap=WORD)

                for x in full_user:
                    words = x.split(" ")

                    for word in words:
                        if len(word) > 1:
                            word0 = word[0]
                            word1 = word[-1]

                            if not word0.isalnum():
                                word = word.replace(word0, "")
                            if not word1.isalnum():
                                word = word.replace(word1, "")

                            tag = ""
                            if word in ambiguous_list:
                                tag = "amb"
                            elif word in adjective_list:
                                tag = "adj"
                            user_canvas.insert(INSERT, word, tag)
                        else:
                            user_canvas.insert(INSERT, word)

                        if word != "\n":
                            user_canvas.insert(INSERT, " ")

                user_canvas.tag_configure("amb", font=Font(size=13, weight="bold"), foreground="dodgerblue4")
                user_canvas.tag_configure("adj", font=Font(size=13, weight="bold"), foreground="mediumpurple3")

                user_canvas.pack()
                user_canvas.place(x=25, y=251)
                user_canvas.config(state=DISABLED)

                srs_canvas = ScrolledText(win, width=70, wrap=WORD)
                for x in full_srs:
                    srs_canvas.insert(INSERT, x)
                srs_canvas.pack()
                srs_canvas.place(x=675, y=251)
                srs_canvas.config(state=DISABLED)

                user_doc_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                                       fg='black', text="User Document")
                user_doc_label.pack()
                user_doc_label.place(x=240, y=210)

                srs_doc_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                                      fg='black', text="SRS Document")
                srs_doc_label.pack()
                srs_doc_label.place(x=915, y=210)
    else:
        showinfo("Error", "Please log in first")

def show_signup_form():
    global or_label, username_label, username_entry, password_label, password_entry, sign_up_button, login_menu
    global password_entry_2, password_label_2, current_screen

    def validate_signup(uname, pwd, pwd2):
        global is_logged_in, upload_menu, or_label

        username_input = uname.get().lower()
        pwd_input = pwd.get()
        pwd2_input = pwd2.get()

        user_list = pandas.read_csv("user.csv").to_dict(orient="records")
        username_list = [x["username"].lower() for x in user_list]
        is_username_exist = username_input in username_list

        if is_username_exist:
            messagebox.showerror("Error", "Username already exists")
        elif len(username_input) > 0 and len(pwd_input) > 0 and len(pwd2_input) > 0 and pwd_input == pwd2_input:
            new_user = {
                "username": uname.get(),
                "password": pwd.get()
            }
            user_list.append(new_user)
            pandas.DataFrame(user_list).to_csv("user.csv", index=False)
            is_logged_in = True
            showinfo("Success", f"Sign up success. Welcome {uname.get()}")

            show_user_icon(uname.get())
            login_menu.destroy()
            sign_up_menu_button.destroy()
            destroy_signup_form()

            uname.set("")
            pwd.set("")
            pwd2.set("")

        if is_logged_in:
            destroy_login_form()
            login_menu.destroy()
            sign_up_menu_button.destroy()
            upload_menu = Button(win, font=("Californian FB", 11, 'bold '), bg='SteelBlue2', fg='Midnight Blue',
                                 width=12,
                                 text="Upload", activebackground='SlateBlue2', activeforeground='Midnight Blue',
                                 command=lambda: show_upload_form())
            upload_menu.pack()
            upload_menu.place(x=575, y=75)
            or_label.destroy()

    if current_screen != "signup":
        current_screen = "signup"

        destroy_upload_form()
        destroy_login_form()
        destroy_review_form()
        destroy_match_form()

        username = StringVar()
        password = StringVar()
        password_2 = StringVar()
        

        username_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                               fg='Midnight Blue',
                               text="Enter User Name")
        username_label.pack()
        username_label.place(x=578, y=160)

        username_entry = Entry(win, textvariable=username)
        username_entry.pack()
        username_entry.place(x=585, y=207)
        username_entry.focus()

        password_label = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                               fg='Midnight Blue',
                               text="Enter Password")
        password_label.pack()
        password_label.place(x=583, y=247)

        password_entry = Entry(win, textvariable=password, show='*')
        password_entry.pack()
        password_entry.place(x=585, y=290)

        password_label_2 = Label(win, font=("Californian FB", 14, 'bold', "underline"), bg='Light Sky Blue',
                                 fg='Midnight Blue',
                                 text="Confirm Password")
        password_label_2.pack()
        password_label_2.place(x=570, y=330)

        password_entry_2 = Entry(win, textvariable=password_2, show='*')
        password_entry_2.pack()
        password_entry_2.place(x=585, y=377)

        validate_signup = partial(validate_signup, username, password, password_2)

        sign_up_button = Button(win, font=("Californian FB", 12, 'bold '), bg='SteelBlue2', fg='Midnight Blue',
                                text="   Sign Up   ", command=validate_signup, activebackground='SlateBlue2',
                                activeforeground='Midnight Blue')
        sign_up_button.pack()
        sign_up_button.place(x=603, y=445)
        
def reset_color():
    login_menu['fg'] = 'Midnight Blue'
    login_menu['bg'] = 'SteelBlue2'

def clickme():
    login_menu['fg'] = 'Midnight Blue'
    login_menu['bg'] = 'SlateBlue2'
    win.after(5000, reset_color)
    
    
def reset_color1():
    sign_up_menu_button['fg'] = 'Midnight Blue'
    sign_up_menu_button['bg'] = 'SteelBlue2'

def clickme1():
    sign_up_menu_button['fg'] = 'Midnight Blue'
    sign_up_menu_button['bg'] = 'SlateBlue2'
    win.after(5000, reset_color1)
    
    
def reset_color2():
    login_menu['fg'] = 'Midnight Blue'
    login_menu['bg'] = 'SteelBlue2'

def clickme2():
    login_menu['fg'] = 'Midnight Blue'
    login_menu['bg'] = 'SlateBlue2'
    win.after(5000, reset_color2)

    
def reset_color3():
    sign_up_menu_button['fg'] = 'Midnight Blue'
    sign_up_menu_button['bg'] = 'SteelBlue2'

def clickme3():
    sign_up_menu_button['fg'] = 'Midnight Blue'
    sign_up_menu_button['bg'] = 'SlateBlue2'
    win.after(5000, reset_color3)

    
def reset_color4():
    upload_menu['fg'] = 'Midnight Blue'
    upload_menu['bg'] = 'SteelBlue2'

def clickme4():
    upload_menu['fg'] = 'Midnight Blue'
    upload_menu['bg'] = 'SlateBlue2'
    win.after(5000, reset_color4)

    
def reset_color5():
    review_menu['fg'] = 'Midnight Blue'
    review_menu['bg'] = 'SteelBlue2'

def clickme5():
    review_menu['fg'] = 'Midnight Blue'
    review_menu['bg'] = 'SlateBlue2'
    win.after(5000, reset_color5)

    
def reset_color6():
    match_menu['fg'] = 'Midnight Blue'
    match_menu['bg'] = 'SteelBlue2'

def clickme6():
    match_menu['fg'] = 'Midnight Blue'
    match_menu['bg'] = 'SlateBlue2'
    win.after(5000, reset_color6)

    
def reset_color7():
    adjective_menu['fg'] = 'Midnight Blue'
    adjective_menu['bg'] = 'SteelBlue2'

def clickme7():
    adjective_menu['fg'] = 'Midnight Blue'
    adjective_menu['bg'] = 'SlateBlue2'
    win.after(5000, reset_color7)

    
def reset_color8():
    compare_heading_menu['fg'] = 'Midnight Blue'
    compare_heading_menu['bg'] = 'SteelBlue2'

def clickme8():
    compare_heading_menu['fg'] = 'Midnight Blue'
    compare_heading_menu['bg'] = 'SlateBlue2'
    win.after(5000, reset_color8)

    
def reset_color9():
    ambiguous_menu['fg'] = 'Midnight Blue'
    ambiguous_menu['bg'] = 'SteelBlue2'

def clickme9():
    ambiguous_menu['fg'] = 'Midnight Blue'
    ambiguous_menu['bg'] = 'SlateBlue2'
    win.after(5000, reset_color9)

    
def reset_color10():
    anaphoric_menu['fg'] = 'Midnight Blue'
    anaphoric_menu['bg'] = 'SteelBlue2'

def clickme10():
    anaphoric_menu['fg'] = 'Midnight Blue'
    anaphoric_menu['bg'] = 'SlateBlue2'
    win.after(5000, reset_color10)

    
def reset_color11():
    cordination_menu['fg'] = 'Midnight Blue'
    cordination_menu['bg'] = 'SteelBlue2'

def clickme11():
    cordination_menu['fg'] = 'Midnight Blue'
    cordination_menu['bg'] = 'SlateBlue2'
    win.after(5000, reset_color11)

    
def reset_color12():
    edit_ambiguous_word_button['fg'] = 'Midnight Blue'
    edit_ambiguous_word_button['bg'] = 'SteelBlue2'
    
def clickme12():
    edit_ambiguous_word_button['fg'] = 'Midnight Blue'
    edit_ambiguous_word_button['bg'] = 'SlateBlue2'
    win.after(5000, reset_color12)
    
    
def reset_color13():
    ambiguous_user_doc_button['fg'] = 'Midnight Blue'
    ambiguous_user_doc_button['bg'] = 'SteelBlue2'
    
def clickme13():
    ambiguous_user_doc_button['fg'] = 'Midnight Blue'
    ambiguous_user_doc_button['bg'] = 'SlateBlue2'
    win.after(5000, reset_color13)
    
def reset_color14():
    excessive_menu['fg'] = 'Midnight Blue'
    excessive_menu['bg'] = 'SteelBlue2'
    
def clickme14():
    excessive_menu['fg'] = 'Midnight Blue'
    excessive_menu['bg'] = 'SlateBlue2'
    win.after(5000, reset_color14)
    
def reset_color15():
    completeness_menu['fg'] = 'Midnight Blue'
    completeness_menu['bg'] = 'SteelBlue2'
    
def clickme15():
    completeness_menu['fg'] = 'Midnight Blue'
    completeness_menu['bg'] = 'SlateBlue2'
    win.after(5000, reset_color15)        

win = Tk()
win.title('Natural Language Processing')
win.minsize(width=1280, height=765)
win.config(bg='Light Sky Blue')

frame = LabelFrame(
    win,
    text='Intelligent Analysis of Software Requirements',
    font=("Californian FB", "28", "bold"),
    bg='Light Sky Blue',
    fg='Midnight Blue'
)
frame.pack(expand=True, fill=BOTH)

login_menu = Button(win, font=("Californian FB", 11, 'bold '), bg='SteelBlue2', fg='Midnight Blue', width=14,
                    text="Login",
                    activebackground='SlateBlue2', activeforeground='Midnight Blue',
                    command=lambda:[show_login_form(), clickme()])
login_menu.pack()
login_menu.place(x=433, y=85)

sign_up_menu_button = Button(win, font=("Californian FB", 11, 'bold '), bg='SteelBlue2', fg='Midnight Blue', width=14,
                             text="Sign Up", activebackground='SlateBlue2',
                             activeforeground='Midnight Blue',
                             command=lambda:[show_signup_form(),clickme1()])
sign_up_menu_button.pack()
sign_up_menu_button.place(x=723, y=85)

or_label = Label(win, font=("Californian FB", 16, 'bold '), bg='Light Sky Blue', fg='black', text="OR")
or_label.pack()
or_label.place(x=627, y=85)

logo_image = ImageTk.PhotoImage(file="admin.png")

win.mainloop()


# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:




